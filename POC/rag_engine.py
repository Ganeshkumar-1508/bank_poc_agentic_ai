# rag_engine.py
# ---------------------------------------------------------------------------
# RAG (Retrieval-Augmented Generation) engine for policy documents.
# Handles document ingestion, chunking, embedding, storage (ChromaDB),
# and retrieval. Used by both the Streamlit upload UI and the CrewAI tool.
# ---------------------------------------------------------------------------

import os
import re
import hashlib
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Optional, Tuple

# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------

# Directory where uploaded raw documents are stored
UPLOAD_DIR = Path(__file__).resolve().parent.parent / "rag_uploads"
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

# ChromaDB persistent storage directory
CHROMA_DIR = Path(__file__).resolve().parent.parent / "rag_chroma_db"
CHROMA_DIR.mkdir(parents=True, exist_ok=True)

# Chunking parameters
CHUNK_SIZE = (
    1500  # characters per chunk (larger = more complete policy rules per chunk)
)
CHUNK_OVERLAP = 300  # overlap between consecutive chunks (reduces boundary cuts)

# ChromaDB collection name
COLLECTION_NAME = "policy_documents"

# Supported file extensions
SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".md", ".docx", ".doc"}


# ---------------------------------------------------------------------------
# Document Loaders
# ---------------------------------------------------------------------------


def load_pdf(file_path: Path) -> str:
    """Extract text from a PDF file with table-aware extraction.

    Tries pdfplumber first (preserves table structure), falls back to pypdf.
    Tables are converted to pipe-delimited rows so chunking and retrieval
    keep related values together.
    """
    # --- Attempt 1: pdfplumber (best table support) ---
    try:
        import pdfplumber

        return _load_pdf_with_pdfplumber(file_path)
    except ImportError:
        pass

    # --- Attempt 2: pypdf fallback (basic, no table parsing) ---
    try:
        from pypdf import PdfReader

        reader = PdfReader(str(file_path))
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text and text.strip():
                pages.append(text.strip())
        return "\n\n".join(pages)
    except Exception:
        # --- Attempt 3: OCR fallback for scanned PDFs ---
        try:
            return _load_pdf_with_ocr(file_path)
        except Exception:
            return ""


def _load_pdf_with_pdfplumber(file_path: Path) -> str:
    """Extract text and tables from PDF using pdfplumber.

    Tables are rendered as pipe-delimited rows with a header prefix so
    downstream chunking treats each row as an atomic unit.
    """
    import pdfplumber

    pages = []
    with pdfplumber.open(str(file_path)) as pdf:
        for page in pdf.pages:
            page_parts = []

            # 1) Body text (outside tables)
            text = page.extract_text() or ""
            if text and text.strip():
                page_parts.append(text.strip())

            # 2) Tables — each table becomes a block of pipe-delimited rows
            tables = page.extract_tables()
            for table_idx, table in enumerate(tables):
                if not table:
                    continue
                table_lines = [f"[TABLE {table_idx + 1}]"]
                for row in table:
                    # Clean None values, strip whitespace
                    cleaned = [str(cell).strip() if cell else "" for cell in row]
                    row_text = " | ".join(cleaned)
                    if row_text.replace(" |", "").replace("| ", "").strip():
                        table_lines.append(row_text)
                if len(table_lines) > 1:  # has at least a header + data
                    page_parts.append("\n".join(table_lines))

            combined = "\n\n".join(page_parts)
            if combined.strip():
                pages.append(combined.strip())

    return "\n\n".join(pages)


def _load_pdf_with_ocr(file_path: Path) -> str:
    """OCR fallback for scanned/image-based PDFs using pytesseract."""
    try:
        from pdf2image import convert_from_path
        import pytesseract

        images = convert_from_path(str(file_path), dpi=300)
        pages = []
        for img in images:
            text = pytesseract.image_to_string(img)
            if text and text.strip():
                pages.append(text.strip())
        return "\n\n".join(pages)
    except ImportError:
        return ""


def load_txt(file_path: Path) -> str:
    """Read text from a plain text or markdown file."""
    encodings = ["utf-8", "utf-8-sig", "latin-1", "iso-8859-1"]
    for enc in encodings:
        try:
            return file_path.read_text(encoding=enc)
        except (UnicodeDecodeError, UnicodeError):
            continue
    raise ValueError(f"Could not decode {file_path.name} with any supported encoding.")


def load_docx(file_path: Path) -> str:
    """Extract text and tables from a .docx file using python-docx.

    Tables are rendered as pipe-delimited rows with a [TABLE N] header,
    matching the PDF table format so chunking handles them consistently.
    """
    from docx import Document

    doc = Document(str(file_path))
    paragraphs = []

    # Body paragraphs
    for para in doc.paragraphs:
        if para.text and para.text.strip():
            paragraphs.append(para.text.strip())

    # Tables — pipe-delimited rows
    for table_idx, table in enumerate(doc.tables):
        table_lines = [f"[TABLE {table_idx + 1}]"]
        for row in table.rows:
            # Deduplicate merged cells (python-docx repeats merged cell text)
            seen_texts = []
            for cell in row.cells:
                cell_text = cell.text.strip() if cell.text else ""
                if cell_text and cell_text not in seen_texts:
                    seen_texts.append(cell_text)
            row_text = " | ".join(seen_texts)
            if row_text.strip():
                table_lines.append(row_text)
        if len(table_lines) > 1:
            paragraphs.append("\n".join(table_lines))

    return "\n\n".join(paragraphs)


def load_document(file_path: Path) -> str:
    """Load a document based on its file extension."""
    ext = file_path.suffix.lower()
    if ext == ".pdf":
        return load_pdf(file_path)
    elif ext in (".txt", ".md"):
        return load_txt(file_path)
    elif ext in (".docx", ".doc"):
        return load_docx(file_path)
    else:
        raise ValueError(f"Unsupported file format: {ext}")


# ---------------------------------------------------------------------------
# Text Chunking
# ---------------------------------------------------------------------------

# Regex to detect table blocks produced by load_pdf / load_docx
_TABLE_BLOCK_RE = re.compile(
    r"(\[TABLE \d+\]\n(?:.*\n)*?)(?=\[TABLE \d+\]|\Z)", re.DOTALL
)


def chunk_text(
    text: str, chunk_size: int = CHUNK_SIZE, chunk_overlap: int = CHUNK_OVERLAP
) -> List[str]:
    """
    Split text into overlapping chunks of approximately `chunk_size` characters.

    Table-aware: detects `[TABLE N]` blocks and keeps table rows together.
    Splits between table rows (at newlines within a table block) rather than
    cutting through the middle of a row.
    """
    if not text or not text.strip():
        return []

    # Normalize whitespace but preserve single newlines inside tables
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = text.strip()

    if len(text) <= chunk_size:
        return [text]

    chunks = []
    start = 0

    while start < len(text):
        end = start + chunk_size

        if end < len(text):
            # --- Check if we're inside a [TABLE ...] block ---
            table_region = _find_table_block(text, start, end)

            if table_region:
                # We're cutting through a table. Find the nearest row boundary.
                row_break = text.rfind("\n", int(start + chunk_size * 0.3), end)
                if row_break > start:
                    end = row_break + 1
                # else: table row is too long, allow the cut (unavoidable)
            else:
                # Normal text — use standard boundary detection
                paragraph_break = text.rfind("\n\n", start, end)
                if paragraph_break > int(start + chunk_size * 0.3):
                    end = paragraph_break + 2
                else:
                    sentence_break = text.rfind(". ", start, end)
                    if sentence_break > int(start + chunk_size * 0.3):
                        end = sentence_break + 2
                    else:
                        line_break = text.rfind("\n", start, end)
                        if line_break > int(start + chunk_size * 0.3):
                            end = line_break + 1

        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)

        start = end - chunk_overlap
        if start < len(text) and start >= end:
            start = end

    return chunks


def _find_table_block(text: str, start: int, end: int) -> Optional[re.Match]:
    """Check if the text region [start:end] overlaps with any [TABLE N] block.
    Returns the Match object if found, None otherwise.
    """
    for m in _TABLE_BLOCK_RE.finditer(text):
        block_start = m.start()
        block_end = m.end()
        # Overlap if the chunk window intersects the table block
        if start < block_end and end > block_start:
            return m
    return None


# ---------------------------------------------------------------------------
# ChromaDB Client (lazy singleton)
# ---------------------------------------------------------------------------

_chroma_client = None
_collection = None


def _get_chroma_client():
    """Get or create the ChromaDB persistent client."""
    global _chroma_client
    if _chroma_client is None:
        import chromadb

        _chroma_client = chromadb.PersistentClient(path=str(CHROMA_DIR))
    return _chroma_client


def _get_collection():
    """Get or create the policy documents collection."""
    global _collection
    if _collection is None:
        client = _get_chroma_client()
        _collection = client.get_or_create_collection(
            name=COLLECTION_NAME,
            metadata={"description": "Banking policy documents for RAG"},
        )
    return _collection


# ---------------------------------------------------------------------------
# Document Ingestion
# ---------------------------------------------------------------------------


def _file_hash(file_path: Path) -> str:
    """Compute SHA-256 hash of a file for deduplication."""
    h = hashlib.sha256()
    with open(file_path, "rb") as f:
        for block in iter(lambda: f.read(8192), b""):
            h.update(block)
    return h.hexdigest()


def ingest_document(
    file_path: Path, category: str = "general", tags: List[str] = None
) -> Dict:
    """Ingest a document into the RAG system (file path based)."""
    collection = _get_collection()
    file_hash = _file_hash(file_path)

    # Check if already ingested
    existing = collection.get(where={"file_hash": file_hash}, limit=1)
    if existing and existing["ids"]:
        return {
            "file_name": file_path.name,
            "file_hash": file_hash,
            "total_chunks": len(existing["ids"]),
            "category": category,
            "status": "skipped (already exists)",
        }

    try:
        text = load_document(file_path)
    except Exception as e:
        return {
            "file_name": file_path.name,
            "file_hash": file_hash,
            "total_chunks": 0,
            "category": category,
            "status": f"error: {str(e)}",
        }

    if not text or len(text.strip()) < 20:
        return {
            "file_name": file_path.name,
            "file_hash": file_hash,
            "total_chunks": 0,
            "category": category,
            "status": "error: document has too little text content",
        }

    chunks = chunk_text(text)
    if not chunks:
        return {
            "file_name": file_path.name,
            "file_hash": file_hash,
            "total_chunks": 0,
            "category": category,
            "status": "error: no chunks created",
        }

    ids, documents, metadatas = [], [], []
    for i, chunk in enumerate(chunks):
        ids.append(f"{file_hash}_{i}")
        documents.append(chunk)
        metadatas.append(
            {
                "file_name": file_path.name,
                "file_hash": file_hash,
                "category": category,
                "tags": ",".join(tags) if tags else "",
                "chunk_index": i,
                "total_chunks": len(chunks),
                "ingested_at": datetime.now().isoformat(),
            }
        )

    batch_size = 100
    for start_idx in range(0, len(ids), batch_size):
        collection.add(
            ids=ids[start_idx : start_idx + batch_size],
            documents=documents[start_idx : start_idx + batch_size],
            metadatas=metadatas[start_idx : start_idx + batch_size],
        )

    return {
        "file_name": file_path.name,
        "file_hash": file_hash,
        "total_chunks": len(chunks),
        "category": category,
        "status": "success",
    }


def ingest_from_bytes(
    file_bytes: bytes, file_name: str, category: str = "general", tags: List[str] = None
) -> Dict:
    """Ingest a document from raw bytes (e.g., from Streamlit file uploader)."""
    ext = Path(file_name).suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        return {
            "file_name": file_name,
            "file_hash": "",
            "total_chunks": 0,
            "category": category,
            "status": f"error: unsupported format '{ext}'",
        }
    file_path = UPLOAD_DIR / file_name
    file_path.write_bytes(file_bytes)
    return ingest_document(file_path, category=category, tags=tags)


# ---------------------------------------------------------------------------
# Retrieval
# ---------------------------------------------------------------------------


def retrieve(
    query: str, n_results: int = 5, category: str = None, tags: List[str] = None
) -> List[Dict]:
    """Retrieve relevant document chunks for a given query."""
    collection = _get_collection()
    where_filter = None
    conditions = []
    if category:
        conditions.append({"category": category})
    if tags:
        tag_conditions = [{"tags": {"$contains": tag}} for tag in tags]
        if len(tag_conditions) == 1:
            conditions.append(tag_conditions[0])
        elif len(tag_conditions) > 1:
            conditions.append({"$or": tag_conditions})
    if len(conditions) == 1:
        where_filter = conditions[0]
    elif len(conditions) > 1:
        where_filter = {"$and": conditions}

    kwargs = {
        "query_texts": [query],
        "n_results": min(n_results, collection.count() or 1),
    }
    if where_filter:
        kwargs["where"] = where_filter

    try:
        results = collection.query(**kwargs)
    except Exception as e:
        return [
            {
                "content": f"RAG retrieval error: {str(e)}",
                "file_name": "",
                "category": "",
                "chunk_index": -1,
                "relevance_score": 0,
                "distance": 1.0,
            }
        ]

    formatted = []
    if results and results["documents"] and results["documents"][0]:
        for i, doc in enumerate(results["documents"][0]):
            metadata = (
                results["metadatas"][0][i]
                if results["metadatas"] and results["metadatas"][0]
                else {}
            )
            distance = (
                results["distances"][0][i]
                if results["distances"] and results["distances"][0]
                else 1.0
            )
            relevance_score = max(0, 1 - distance)
            formatted.append(
                {
                    "content": doc,
                    "file_name": metadata.get("file_name", "unknown"),
                    "category": metadata.get("category", ""),
                    "tags": metadata.get("tags", ""),
                    "chunk_index": metadata.get("chunk_index", i),
                    "relevance_score": round(relevance_score, 4),
                    "distance": round(distance, 4),
                }
            )
    return formatted


def retrieve_as_text(
    query: str, n_results: int = 5, category: str = None, tags: List[str] = None
) -> str:
    """Retrieve relevant chunks formatted for LLM prompts."""
    results = retrieve(query, n_results=n_results, category=category, tags=tags)
    if not results:
        return "No relevant policy documents found for this query."
    output_parts = [f"=== RELEVANT POLICY DOCUMENTS (top {len(results)} results) ===\n"]
    for i, r in enumerate(results, 1):
        output_parts.append(
            f"--- Result {i} ---\nSource: {r['file_name']} | "
            f"Category: {r['category']} | Relevance: {r['relevance_score']:.1%}\n"
            f"{r['content']}\n"
        )
    output_parts.append("=== END OF POLICY DOCUMENTS ===")
    return "\n".join(output_parts)


# ---------------------------------------------------------------------------
# Document Management
# ---------------------------------------------------------------------------


def list_documents() -> List[Dict]:
    """List all unique documents in the RAG system."""
    collection = _get_collection()
    if collection.count() == 0:
        return []
    all_data = collection.get()
    if not all_data or not all_data["metadatas"]:
        return []
    docs = {}
    for metadata in all_data["metadatas"]:
        fh = metadata.get("file_hash", "")
        if fh and fh not in docs:
            docs[fh] = {
                "file_name": metadata.get("file_name", "unknown"),
                "file_hash": fh,
                "category": metadata.get("category", ""),
                "tags": metadata.get("tags", ""),
                "total_chunks": metadata.get("total_chunks", 0),
                "ingested_at": metadata.get("ingested_at", ""),
            }
    return list(docs.values())


def delete_document(file_hash: str) -> Dict:
    collection = _get_collection()
    collection.delete(where={"file_hash": file_hash})
    return {"file_hash": file_hash, "status": "deleted"}


def delete_document_by_name(file_name: str) -> Dict:
    collection = _get_collection()
    results = collection.get(where={"file_name": file_name}, limit=1)
    if not results or not results["metadatas"]:
        return {"file_name": file_name, "status": "not found"}
    return delete_document(results["metadatas"][0].get("file_hash", ""))


def get_stats() -> Dict:
    collection = _get_collection()
    docs = list_documents()
    return {
        "total_chunks": collection.count(),
        "total_documents": len(docs),
        "categories": list(set(d["category"] for d in docs if d.get("category"))),
        "storage_dir": str(CHROMA_DIR),
        "upload_dir": str(UPLOAD_DIR),
    }


def reset_rag():
    global _chroma_client, _collection
    _chroma_client = None
    _collection = None
    import shutil

    if CHROMA_DIR.exists():
        shutil.rmtree(CHROMA_DIR)
    CHROMA_DIR.mkdir(parents=True, exist_ok=True)
    return {"status": "RAG system reset complete"}
    
    
# ---------------------------------------------------------------------------
# Query Interface (for CrewAI integration)
# ---------------------------------------------------------------------------

def query_rag(query: str, n_results: int = 5, category: str = None, tags: List[str] = None) -> str:
    """Query the RAG system and return formatted results.

    This is the main entry point for CrewAI agents to query policy documents.
    Returns formatted text suitable for LLM prompts.

    Args:
        query: The query text to search for
        n_results: Number of results to return (default: 5)
        category: Optional category filter
        tags: Optional list of tag filters

    Returns:
        Formatted string with relevant policy documents
    """
    return retrieve_as_text(query, n_results=n_results, category=category, tags=tags)
